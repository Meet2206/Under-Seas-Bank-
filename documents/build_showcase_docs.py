from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(85, 85, 85)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)


def style_document(doc, title, subtitle):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True

    header = section.header.paragraphs[0]
    header.text = "Underseas Bank | Showcase Documentation"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = GRAY

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(23)
    r.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(subtitle)
    r.font.size = Pt(13)
    r.font.color.rgb = GRAY


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, "F2F4F7")
        set_cell_text(cell, header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_code_doc(path):
    doc = Document()
    style_document(
        doc,
        "Code Explanation and Import Guide",
        "Module-by-module explanation of the banking system, its imports, and why each part matters.",
    )

    doc.add_heading("1. System Overview", level=1)
    doc.add_paragraph(
        "The project is split into a FastAPI backend and a React/Vite frontend. "
        "The backend handles authentication, banking operations, persistence, OTP/email flows, and analytics. "
        "The frontend renders the user interface and calls the backend through a central API service."
    )

    doc.add_heading("2. Backend Entry and Configuration", level=1)
    add_table(doc, ["File", "Important imports", "Why it matters"], [
        ("app/main.py", "FastAPI, CORSMiddleware, routers, database Base/engine", "Starts the API, creates tables on startup, exposes health endpoints, and registers all route groups."),
        ("app/config.py", "BaseSettings, lru_cache", "Loads environment variables once and centralizes database, JWT, Redis, SMTP, OTP, and encryption settings."),
        ("app/database.py", "create_engine, declarative_base, sessionmaker", "Creates the SQLAlchemy engine and provides request-scoped database sessions."),
        ("app/redis_client.py", "redis, Optional", "Defines Redis helpers for OTP/cache/rate-limiting style use cases."),
    ])

    doc.add_heading("3. Runtime Request Flow", level=1)
    add_bullets(doc, [
        "A browser action in React calls a function from `frontend/src/services/api.js`.",
        "The request reaches a FastAPI router such as `/transactions/transfer` or `/accounts/my-accounts`.",
        "Protected routes call `get_current_user`, which validates the bearer token and loads the authenticated user.",
        "The route validates ownership, then delegates business logic to a service module.",
        "The service reads or writes SQLAlchemy models through a request-scoped database session.",
        "The response is serialized back to JSON and rendered by the React page that initiated the action.",
    ])

    doc.add_heading("4. Data Models", level=1)
    add_table(doc, ["Model", "Stored data", "Importance"], [
        ("User", "name, email, phone number, MPIN hash, lock/verification flags", "Represents customers and login state."),
        ("Account", "encrypted account number, type, owner, balance", "Represents deposit accounts and balances."),
        ("Transaction", "source account, destination account, amount, type, timestamp", "Stores the financial ledger events."),
        ("Loan", "principal, rate, tenure, EMI, remaining balance", "Tracks borrowing products."),
        ("FixedDeposit", "principal, rate, term, maturity amount, status", "Tracks term deposits."),
        ("CreditCard", "encrypted card number, limits, usage, status", "Tracks card products and utilization."),
        ("Beneficiary", "recipient identity and encrypted account number", "Stores saved transfer recipients."),
    ])

    doc.add_heading("5. Database Design Details", level=1)
    add_table(doc, ["Table", "Primary relationships", "Design notes"], [
        ("users", "One user to many accounts/products", "Stores hashed MPINs and verification flags; user identity is the ownership anchor for the rest of the system."),
        ("accounts", "Belongs to users", "Stores encrypted account numbers plus keyed lookup hashes for safe exact-match uniqueness."),
        ("transactions", "References source and destination accounts", "Acts as the financial event log used by statements, passbooks, and analytics."),
        ("credit_cards", "Belongs to user and linked account", "Stores encrypted card numbers and available/used credit values."),
        ("loans", "Belongs to user and linked account", "Stores amount, tenure, EMI, and remaining balance."),
        ("fixed_deposits", "Belongs to user and linked account", "Stores principal, rate, duration, maturity value, and status."),
        ("beneficiaries", "Belongs to user", "Stores saved recipient information for transfer convenience."),
    ])

    doc.add_heading("6. Schemas and Validation", level=1)
    doc.add_paragraph(
        "The files in `app/schemas` use Pydantic models to validate incoming requests and define response shape. "
        "For example, `auth_schema.py` restricts MPINs to four digits, phone numbers to ten digits, and OTPs to six digits. "
        "This prevents many malformed inputs before business logic runs."
    )

    doc.add_heading("7. Services", level=1)
    add_table(doc, ["Service", "Key imports", "Role"], [
        ("auth_service.py", "HTTPException, Session, IntegrityError, User, password_hash, jwt_handler", "Registers users, hashes MPINs, handles login, lockout, and token creation."),
        ("transaction_service.py", "HTTPException, Session, Account, Transaction", "Implements deposits, withdrawals, transfers, and transaction lookup."),
        ("account_service.py", "secrets, Account, Transaction", "Creates account numbers and generates statements/passbooks."),
        ("loan_service.py", "HTTPException, Loan, Account, calculate_emi", "Creates loans and accepts EMI payments after ownership checks."),
        ("fd_service.py", "HTTPException, FixedDeposit, Account", "Creates and closes fixed deposits after ownership checks."),
        ("creditcard_service.py", "HTTPException, CreditCard, Account", "Issues cards and processes purchases/payments after ownership checks."),
        ("otp_service.py", "secrets, time, threading", "Generates secure OTPs, stores them with expiry, verifies them, and sends email OTPs."),
        ("email_service.py", "smtplib, threading, email.mime", "Sends welcome emails and OTP emails through SMTP in background threads."),
    ])

    doc.add_heading("8. Routes and Middleware", level=1)
    add_bullets(doc, [
        "`auth_routes.py` exposes register, login, profile, email OTP, phone OTP, and email status endpoints.",
        "`transaction_routes.py` protects money movement endpoints and checks that the source account belongs to the logged-in user.",
        "`account_routes.py` exposes account creation, statements, and passbooks with ownership checks.",
        "`loan_routes.py`, `fd_routes.py`, and `creditcard_routes.py` expose product workflows.",
        "`analytics_routes.py` returns expense summaries only for the current user's accounts.",
        "`auth_middleware.py` extracts the bearer token, decodes it, and loads the current user.",
    ])

    doc.add_heading("9. Security-Oriented Utilities", level=1)
    add_bullets(doc, [
        "`password_hash.py` uses bcrypt through Passlib so MPINs are not stored in plaintext.",
        "`jwt_handler.py` signs and verifies access tokens.",
        "`encryption.py` uses Fernet encryption to protect account numbers, beneficiary account numbers, and card numbers at rest.",
        "`otp_service.py` now uses `secrets` instead of predictable pseudo-random generation.",
    ])

    doc.add_heading("10. Frontend Structure", level=1)
    add_table(doc, ["Area", "Purpose"], [
        ("src/App.jsx", "Defines the route map for login, dashboard, accounts, transfers, products, and analytics."),
        ("src/services/api.js", "Centralizes HTTP requests to the backend and attaches bearer tokens."),
        ("pages/*.jsx", "Implements each banking screen and user workflow."),
        ("components/*.jsx", "Reusable visual pieces such as sidebar, header, cards, charts, and transaction tables."),
        ("layout/MainLayout.jsx", "Wraps authenticated pages with shared navigation and header structure."),
    ])

    doc.add_heading("11. Endpoint Catalogue", level=1)
    add_table(doc, ["Group", "Endpoints", "Purpose"], [
        ("Health", "`/`, `/health`", "Confirms that the API process is reachable."),
        ("Authentication", "`/auth/register`, `/auth/login`, `/auth/me`", "Creates users, authenticates users, and returns the current profile."),
        ("Email OTP", "`/auth/send-email-otp`, `/auth/verify-email`, `/auth/email-status`", "Issues and verifies email OTPs and exposes safe setup status."),
        ("Accounts", "`/accounts/create`, `/accounts/my-accounts`, `/accounts/statement/{id}`, `/accounts/passbook/{id}`", "Creates and reads account information."),
        ("Transactions", "`/transactions/deposit`, `/withdraw`, `/transfer`, `/history/{id}`", "Executes and lists money movement."),
        ("Products", "`/loans/*`, `/fd/*`, `/credit-card/*`", "Implements loan, fixed-deposit, and card workflows."),
        ("Analytics", "`/analytics/expenses`", "Returns summarized spending data for the authenticated user's accounts."),
        ("Beneficiaries", "`/beneficiaries/add`, `/beneficiaries/my`", "Stores and lists saved recipients."),
    ])

    doc.add_heading("12. Why the Imports Matter", level=1)
    add_bullets(doc, [
        "Framework imports such as FastAPI, React, and SQLAlchemy provide the core application runtime.",
        "Validation imports such as Pydantic models prevent invalid requests from reaching business logic.",
        "Security imports such as bcrypt, jose/jwt, secrets, and Fernet protect credentials, tokens, OTPs, and sensitive identifiers.",
        "Infrastructure imports such as Redis and SMTP libraries connect the app to external services.",
        "Visualization imports such as Recharts convert transaction summaries into readable charts for users.",
    ])

    doc.add_heading("13. Configuration Reference", level=1)
    add_table(doc, ["Setting", "Purpose"], [
        ("DATABASE_URL", "Connects SQLAlchemy to PostgreSQL."),
        ("SECRET_KEY", "Signs JWT access tokens; must stay private and random."),
        ("REDIS_URL", "Provides optional OTP/cache backing store."),
        ("SMTP_HOST / SMTP_PORT / SMTP_USER / SMTP_PASSWORD / EMAIL_FROM", "Enable Gmail SMTP delivery for welcome emails and OTP emails."),
        ("FIELD_ENCRYPTION_KEY", "Encrypts account/card/beneficiary identifiers at rest."),
        ("CORS_ORIGINS", "Lists allowed frontend origins."),
        ("DEBUG", "Controls development-only behavior such as docs exposure."),
    ])

    doc.add_heading("14. Module-by-Module Reading Guide", level=1)
    add_bullets(doc, [
        "Start with `main.py` to understand app startup and registered routes.",
        "Read `config.py` and `.env` next to understand external dependencies.",
        "Read the `models` folder before services so the domain objects are clear.",
        "Read `schemas` to see every accepted request shape and validation rule.",
        "Read `routes` to understand the public API surface.",
        "Read `services` last to follow the business logic behind each endpoint.",
        "On the frontend, start from `App.jsx`, then `services/api.js`, then the page components.",
    ])

    doc.add_heading("15. Current Engineering Strengths", level=1)
    add_bullets(doc, [
        "Clear separation between routers, services, schemas, models, and utilities.",
        "Centralized API calls on the frontend.",
        "Password hashing, JWT authentication, email OTP flow, encrypted sensitive identifiers, and ownership checks are present.",
        "Frontend is organized around reusable layout and component files.",
    ])

    doc.add_heading("16. Areas for Future Growth", level=1)
    add_bullets(doc, [
        "Add a formal Alembic migration history instead of rebuilding schema from metadata.",
        "Expand automated tests across authentication, authorization, transaction rollback, and product workflows.",
        "Add structured audit logs for privileged or high-risk actions.",
        "Introduce refresh tokens and stronger session lifecycle controls if the app evolves beyond a showcase.",
    ])

    doc.save(path)


def build_logic_doc(path):
    doc = Document()
    style_document(
        doc,
        "Logic and Algorithms Used",
        "Plain-language explanation of the major workflows, formulas, and decision logic in the banking system.",
    )

    doc.add_heading("1. Authentication Logic", level=1)
    add_bullets(doc, [
        "Registration checks whether the email or phone number already exists.",
        "The four-digit MPIN is hashed with bcrypt before storage.",
        "Login searches by phone number, verifies the MPIN hash, increments failed attempts on mistakes, and locks the account after five failed attempts.",
        "Successful login resets the failed-attempt counter and returns a signed JWT access token.",
    ])

    doc.add_heading("2. Email OTP Algorithm", level=1)
    add_bullets(doc, [
        "Generate a six-digit code using a cryptographically secure random source.",
        "Store the OTP with a five-minute expiry, preferring Redis and falling back to in-memory storage if Redis is unavailable.",
        "On verification, compare the submitted code, reject expired or mismatched values, and delete the OTP after success to prevent reuse.",
    ])

    doc.add_heading("3. Authorization Algorithm", level=1)
    add_bullets(doc, [
        "Extract the bearer token from the request.",
        "Decode the JWT signature and expiry; reject missing, invalid, or expired tokens.",
        "Load the user identified by the token payload.",
        "For resource-specific actions, check that the target account, loan, deposit, or card belongs to that user.",
        "Only after ownership is confirmed does the service execute the requested banking operation.",
    ])

    doc.add_heading("4. Money Movement Logic", level=1)
    add_table(doc, ["Operation", "Algorithm"], [
        ("Deposit", "Validate positive amount -> confirm account exists -> increase balance -> create deposit transaction -> commit."),
        ("Withdraw", "Validate positive amount -> confirm account exists -> ensure sufficient balance -> decrease balance -> create withdrawal transaction -> commit."),
        ("Transfer", "Validate amount -> load source and destination accounts -> ensure source ownership and sufficient balance -> debit source -> credit destination -> create transfer transaction -> commit atomically."),
    ])

    doc.add_heading("5. Transaction Consistency Rules", level=1)
    add_bullets(doc, [
        "Amounts must be positive before any balance change is accepted.",
        "Withdrawals and transfers reject insufficient balance.",
        "Transfers debit and credit within the same database transaction before commit.",
        "If a failure occurs during deposit, withdrawal, or transfer, the service rolls back the session.",
        "Statements and passbooks are derived from the transaction history rather than manually entered balances.",
    ])

    doc.add_heading("6. Account Statement and Passbook", level=1)
    doc.add_paragraph(
        "Both features load all transactions related to one account and replay them in order. "
        "Deposits add to the running balance, withdrawals subtract from it, and transfers add or subtract depending on whether the account is sender or receiver."
    )

    doc.add_heading("7. Loan EMI Formula", level=1)
    doc.add_paragraph(
        "The EMI calculator uses the standard reducing-balance formula: EMI = P * r * (1 + r)^n / ((1 + r)^n - 1), "
        "where P is principal, r is monthly interest rate, and n is the number of monthly installments."
    )

    doc.add_heading("8. Fixed Deposit Maturity Logic", level=1)
    doc.add_paragraph(
        "The fixed deposit module uses simple-interest maturity logic: maturity = principal * (1 + annual_rate * years). "
        "The result is rounded to two decimal places."
    )

    doc.add_heading("9. Credit Card Logic", level=1)
    add_bullets(doc, [
        "Issuance creates a card linked to a user-owned account with a chosen credit limit.",
        "Purchase first checks available credit, then increases used credit and decreases available credit.",
        "Bill payment decreases used credit and restores available credit, never allowing used credit to stay negative.",
    ])

    doc.add_heading("10. Analytics Logic", level=1)
    doc.add_paragraph(
        "Analytics collects transactions only from the logged-in user's accounts, groups expenses into categories, "
        "and sends summarized values to the frontend for chart rendering."
    )

    doc.add_heading("11. Sensitive-Data Protection Logic", level=1)
    add_bullets(doc, [
        "MPINs are stored as bcrypt hashes, never plaintext.",
        "Account numbers, beneficiary account numbers, and card numbers are encrypted before database storage.",
        "Because encrypted values are randomized, the system also stores keyed lookup hashes for exact-match uniqueness without exposing the original number.",
        "OTP expiry and deletion after success reduce replay risk.",
    ])

    doc.add_heading("12. Error Handling and Failure Behavior", level=1)
    add_table(doc, ["Condition", "System response"], [
        ("Missing token", "Return unauthorized response before loading user data."),
        ("Expired or invalid token", "Reject the request and require login again."),
        ("Negative or zero amount", "Reject request at validation boundary."),
        ("Insufficient account balance", "Reject withdrawal or transfer."),
        ("Foreign resource access", "Reject with unauthorized/forbidden response."),
        ("Invalid or expired OTP", "Reject verification and do not mark the address verified."),
    ])

    doc.add_heading("13. Core Formulas", level=1)
    add_table(doc, ["Feature", "Formula / rule"], [
        ("Loan EMI", "P * r * (1 + r)^n / ((1 + r)^n - 1)"),
        ("FD maturity", "principal * (1 + annual_rate * years)"),
        ("Credit available", "credit_limit - used_credit"),
        ("Running balance", "previous balance adjusted by transaction direction and type"),
    ])

    doc.add_heading("14. Example End-to-End Scenario", level=1)
    add_bullets(doc, [
        "A new user registers with name, email, phone number, and MPIN.",
        "The MPIN is hashed, the user is saved, and a welcome email is queued.",
        "The user requests an email OTP; a six-digit code is generated, stored with expiry, and sent through Gmail SMTP.",
        "After OTP verification, the user logs in and receives a JWT.",
        "The user creates a savings account; the visible account number is decrypted for display but stored encrypted in PostgreSQL.",
        "The user deposits money, transfers money, and later views a statement generated from recorded transactions.",
    ])

    doc.add_heading("15. Recommended Future Improvements", level=1)
    add_bullets(doc, [
        "Add automated tests for authentication, authorization, and transaction consistency.",
        "Move from float amounts to decimal currency handling.",
        "Add database migrations instead of startup-time `create_all`.",
        "Replace phone OTP console logging with a real SMS provider if phone verification is needed.",
        "Add structured logging, audit trails, and stronger production configuration controls.",
    ])

    doc.save(path)


if __name__ == "__main__":
    build_code_doc("documents/Code_Explanation_and_Imports.docx")
    build_logic_doc("documents/Logic_and_Algorithms.docx")
